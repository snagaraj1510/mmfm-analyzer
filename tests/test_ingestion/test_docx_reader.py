"""Tests for DOCX reader."""

from __future__ import annotations

import pytest
from pathlib import Path


class TestDocxReader:
    def test_file_not_found_raises(self):
        from mmfm.ingestion.docx_reader import parse_docx
        with pytest.raises(FileNotFoundError):
            parse_docx(Path("/nonexistent/file.docx"))

    def test_wrong_extension_raises(self, tmp_path):
        from mmfm.ingestion.docx_reader import parse_docx
        bad_file = tmp_path / "test.txt"
        bad_file.write_text("hello")
        with pytest.raises(ValueError, match="Expected .docx"):
            parse_docx(bad_file)

    def test_parses_docx_content(self, tmp_path):
        """Create a real DOCX and verify content extraction."""
        pytest.importorskip("docx")
        from docx import Document

        doc = Document()
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph("This market shows strong investment potential.")
        doc.add_heading("Financial Overview", level=2)
        doc.add_paragraph("NPV is positive at 10% discount rate.")

        path = tmp_path / "test.docx"
        doc.save(str(path))

        from mmfm.ingestion.docx_reader import parse_docx
        result = parse_docx(path)
        assert result.full_text != ""
        assert "investment potential" in result.full_text
        assert result.checksum.startswith("sha256:")
        assert len(result.sections) >= 1
