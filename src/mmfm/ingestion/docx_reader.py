"""
DOCX content extraction using python-docx.

Extracts paragraphs, headings, and table content from Word documents.
Returns structured content suitable for chunking and RAG indexing.
"""

from __future__ import annotations

import hashlib
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional


@dataclass
class DocxSection:
    heading: Optional[str]
    paragraphs: list[str] = field(default_factory=list)


@dataclass
class ParsedDocx:
    source_file: Path
    sections: list[DocxSection] = field(default_factory=list)
    full_text: str = ""
    table_texts: list[str] = field(default_factory=list)
    checksum: str = ""
    metadata: dict = field(default_factory=dict)


def _compute_checksum(path: Path) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return f"sha256:{h.hexdigest()}"


def parse_docx(file_path: Path | str) -> ParsedDocx:
    """
    Parse a .docx file and extract structured text content.

    Args:
        file_path: Path to the .docx file

    Returns:
        ParsedDocx with sections, full text, and table content

    Raises:
        FileNotFoundError: If file does not exist
        ValueError: If file is not a DOCX
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required: pip install python-docx")

    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")
    if path.suffix.lower() not in (".docx", ".doc"):
        raise ValueError(f"Expected .docx file, got: {path.suffix}")

    doc = Document(str(path))

    sections: list[DocxSection] = []
    current_section = DocxSection(heading=None)
    all_text_parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Detect headings by style name
        style_name = para.style.name.lower() if para.style else ""
        is_heading = "heading" in style_name or style_name.startswith("title")

        if is_heading:
            if current_section.paragraphs:
                sections.append(current_section)
            current_section = DocxSection(heading=text)
        else:
            current_section.paragraphs.append(text)
            all_text_parts.append(text)

    if current_section.paragraphs or current_section.heading:
        sections.append(current_section)

    # Extract table content as flat text
    table_texts: list[str] = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            table_text = "\n".join(rows)
            table_texts.append(table_text)
            all_text_parts.append(table_text)

    full_text = "\n\n".join(all_text_parts)

    return ParsedDocx(
        source_file=path,
        sections=sections,
        full_text=full_text,
        table_texts=table_texts,
        checksum=_compute_checksum(path),
        metadata={"filename": path.name, "section_count": len(sections)},
    )
